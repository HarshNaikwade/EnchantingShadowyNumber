import re
import os
import logging
from typing import Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_dates(text: str) -> Tuple[str, str]:
    """Extract effective date and creation date from document text."""
    effective_date = "NA"
    creation_date = "NA"

    date_patterns = [
        r'\b(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\b',
        r'\b(\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December),?\s+\d{4})\b',
        r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b',
    ]

    effective_patterns = [
        r'effective\s+(?:date|from)[:\s]+([^\n,;]+)',
        r'effective\s+as\s+of[:\s]+([^\n,;]+)',
        r'this\s+agreement\s+(?:is\s+)?effective\s+([^\n,;]+)',
    ]

    creation_patterns = [
        r'date\s+of\s+agreement[:\s]+([^\n,;]+)',
        r'dated\s+(?:this\s+)?([^\n,;]+)',
        r'entered\s+into\s+(?:on\s+)?(?:this\s+)?([^\n,;]+)',
        r'executed\s+(?:on\s+)?(?:this\s+)?([^\n,;]+)',
        r'made\s+(?:and\s+entered\s+into\s+)?(?:on\s+)?(?:this\s+)?([^\n,;]+)',
    ]

    text_lower = text.lower()

    for pattern in effective_patterns:
        match = re.search(pattern, text_lower, re.IGNORECASE)
        if match:
            candidate = match.group(1).strip()
            for dp in date_patterns:
                dm = re.search(dp, candidate, re.IGNORECASE)
                if dm:
                    effective_date = dm.group(1)
                    break
            if effective_date != "NA":
                break

    for pattern in creation_patterns:
        match = re.search(pattern, text_lower, re.IGNORECASE)
        if match:
            candidate = match.group(1).strip()
            for dp in date_patterns:
                dm = re.search(dp, candidate, re.IGNORECASE)
                if dm:
                    creation_date = dm.group(1)
                    break
            if creation_date != "NA":
                break

    return effective_date, creation_date


def parse_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber with OCR fallback."""
    text_parts = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"\n--- Page {page_num} ---\n{page_text}")

        full_text = "\n".join(text_parts)
        if len(full_text.strip()) < 100:
            logger.info("Low text content from pdfplumber, trying OCR...")
            return _ocr_pdf(file_path)
        return full_text
    except Exception as e:
        logger.error(f"pdfplumber failed: {e}")
        try:
            return _ocr_pdf(file_path)
        except Exception as ocr_err:
            logger.error(f"OCR also failed: {ocr_err}")
            return ""


def _ocr_pdf(file_path: str) -> str:
    """OCR fallback for PDF files using pdf2image + pytesseract."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image

        images = convert_from_path(file_path, dpi=200)
        text_parts = []
        for i, image in enumerate(images, 1):
            page_text = pytesseract.image_to_string(image)
            text_parts.append(f"\n--- Page {i} (OCR) ---\n{page_text}")
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("OCR dependencies not available (pdf2image/pytesseract)")
        return ""
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return ""


def parse_docx(file_path: str) -> str:
    """Extract text from Word document using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_parts.append(f"Line {i+1}: {para.text}")
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"python-docx failed: {e}")
        return ""


def parse_document(file_path: str, filename: str) -> dict:
    """Main document parsing function. Returns extracted text and metadata."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text = parse_docx(file_path)
    else:
        text = ""
        logger.warning(f"Unsupported file type: {ext}")

    effective_date, creation_date = extract_dates(text)

    return {
        "text": text,
        "effective_date": effective_date,
        "creation_date": creation_date,
        "word_count": len(text.split()) if text else 0,
    }
