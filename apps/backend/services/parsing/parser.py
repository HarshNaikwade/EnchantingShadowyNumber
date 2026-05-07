import re
import logging
from typing import Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_dates(text: str) -> Tuple[str, str]:
    """Extract effective date and creation date from document text."""
    effective_date = "NA"
    creation_date = "NA"

    date_patterns = [
        r'\b(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\b',
        r'\b(\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December),?\s+\d{4})\b',
        r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b',
    ]

    effective_patterns = [
        r'effective\s+(?:date|from)[:\s]+([^\n,;]+)',
        r'effective\s+as\s+of[:\s]+([^\n,;]+)',
        r'this\s+agreement\s+(?:is\s+)?effective\s+([^\n,;]+)',
    ]

    creation_patterns = [
        r'date\s+of\s+agreement[:\s]+([^\n,;]+)',
        r'dated\s+(?:this\s+)?([^\n,;]+)',
        r'entered\s+into\s+(?:on\s+)?(?:this\s+)?([^\n,;]+)',
        r'executed\s+(?:on\s+)?(?:this\s+)?([^\n,;]+)',
        r'made\s+(?:and\s+entered\s+into\s+)?(?:on\s+)?(?:this\s+)?([^\n,;]+)',
    ]

    text_lower = text.lower()

    for pattern in effective_patterns:
        match = re.search(pattern, text_lower, re.IGNORECASE)
        if match:
            candidate = match.group(1).strip()
            for dp in date_patterns:
                dm = re.search(dp, candidate, re.IGNORECASE)
                if dm:
                    effective_date = dm.group(1)
                    break
            if effective_date != "NA":
                break

    for pattern in creation_patterns:
        match = re.search(pattern, text_lower, re.IGNORECASE)
        if match:
            candidate = match.group(1).strip()
            for dp in date_patterns:
                dm = re.search(dp, candidate, re.IGNORECASE)
                if dm:
                    creation_date = dm.group(1)
                    break
            if creation_date != "NA":
                break

    return effective_date, creation_date


def clean_text(text: str) -> str:
    """Normalize whitespace and collapse redundant blank lines."""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------

def parse_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber with OCR fallback."""
    raw, _ = parse_pdf_with_pages(file_path)
    return raw


def parse_pdf_with_pages(file_path: str) -> tuple[str, list[dict]]:
    """
    Parse PDF and return (full_raw_text, per_page_list).
    Each page entry: {page_number, text, char_count}.
    Logs full diagnostic info for debugging.
    """
    try:
        import pdfplumber
        pages_data: list[dict] = []
        text_parts: list[str] = []

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info("PDF PARSE — total pages: %d | file: %s", total_pages, file_path)

            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                pages_data.append({
                    "page_number": page_num,
                    "text": page_text,
                    "char_count": len(page_text),
                })
                if page_text.strip():
                    text_parts.append(f"\n--- Page {page_num} ---\n{page_text}")

        full_text = "\n".join(text_parts)
        logger.info("PDF PARSE — raw_text_length: %d chars", len(full_text))
        logger.info("PDF PARSE — clause_count_estimate: %d (newline-sep paragraphs)", full_text.count("\n\n"))
        logger.info("PDF PARSE — first 1000 chars:\n%s", full_text[:1000])
        logger.info("PDF PARSE — last 1000 chars:\n%s", full_text[-1000:])

        if len(full_text.strip()) < 100:
            logger.info("PDF PARSE — low text content, trying OCR fallback...")
            ocr_text = _ocr_pdf(file_path)
            ocr_pages = [{"page_number": 1, "text": ocr_text, "char_count": len(ocr_text)}]
            return ocr_text, ocr_pages

        return full_text, pages_data

    except Exception as exc:
        logger.error("PDF PARSE — pdfplumber failed: %s", exc)
        try:
            ocr_text = _ocr_pdf(file_path)
            return ocr_text, [{"page_number": 1, "text": ocr_text, "char_count": len(ocr_text)}]
        except Exception as ocr_err:
            logger.error("PDF PARSE — OCR also failed: %s", ocr_err)
            return "", []


def _ocr_pdf(file_path: str) -> str:
    """OCR fallback for PDF files using pdf2image + pytesseract."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(file_path, dpi=200)
        text_parts = []
        for i, image in enumerate(images, 1):
            page_text = pytesseract.image_to_string(image)
            text_parts.append(f"\n--- Page {i} (OCR) ---\n{page_text}")
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("OCR dependencies not available (pdf2image/pytesseract)")
        return ""
    except Exception as exc:
        logger.error("OCR failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------

def parse_docx(file_path: str) -> str:
    """Extract text from Word document using python-docx."""
    raw, _ = parse_docx_with_lines(file_path)
    return raw


def parse_docx_with_lines(file_path: str) -> tuple[str, list[dict]]:
    """
    Parse DOCX and return (full_raw_text, pseudo_pages_list).
    DOCX has no real pages; the whole doc is returned as a single entry.
    Logs full diagnostic info for debugging.
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts: list[str] = []

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_parts.append(f"Line {i + 1}: {para.text}")

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        logger.info("DOCX PARSE — paragraphs: %d | file: %s", len(doc.paragraphs), file_path)
        logger.info("DOCX PARSE — raw_text_length: %d chars", len(full_text))
        logger.info("DOCX PARSE — first 1000 chars:\n%s", full_text[:1000])
        logger.info("DOCX PARSE — last 1000 chars:\n%s", full_text[-1000:])

        pages = [{"page_number": 1, "text": full_text, "char_count": len(full_text)}]
        return full_text, pages

    except Exception as exc:
        logger.error("DOCX PARSE — python-docx failed: %s", exc)
        return "", []


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def parse_document(file_path: str, filename: str) -> dict:
    """Main document parsing function. Returns extracted text and metadata."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text = parse_docx(file_path)
    else:
        text = ""
        logger.warning("Unsupported file type: %s", ext)

    effective_date, creation_date = extract_dates(text)

    return {
        "text": text,
        "effective_date": effective_date,
        "creation_date": creation_date,
        "word_count": len(text.split()) if text else 0,
    }
